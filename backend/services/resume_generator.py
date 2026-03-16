"""Resume generation service — creates DOCX and PDF files from tailored text."""
from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class ResumeGenerator:
    """Generates DOCX and PDF resume files from plain text."""

    def __init__(self, output_dir: str = "./generated_resumes"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_docx(self, resume_text: str, filename_prefix: str = "tailored_resume") -> str:
        """Generate a DOCX file from resume text. Returns the file path."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style the document
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)

            # Split text into sections and add to document
            sections = self._parse_sections(resume_text)
            for section_type, content in sections:
                if section_type == "name":
                    para = doc.add_heading(content, level=1)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif section_type == "header":
                    doc.add_heading(content, level=2)
                elif section_type == "contact":
                    para = doc.add_paragraph(content)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif section_type == "bullet":
                    doc.add_paragraph(content, style="List Bullet")
                else:
                    doc.add_paragraph(content)

            filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = self.output_dir / filename
            doc.save(str(file_path))
            logger.info("Generated DOCX: %s", file_path)
            return str(file_path)

        except Exception as exc:
            logger.error("DOCX generation failed: %s", exc)
            raise

    def generate_pdf(self, resume_text: str, filename_prefix: str = "tailored_resume") -> str:
        """Generate a PDF file from resume text. Returns the file path."""
        try:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)

            # Encode text to handle special characters
            lines = resume_text.split("\n")
            in_header = True

            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(3)
                    continue

                # Detect section headers (ALL CAPS or ends with colon)
                if (line.isupper() and len(line) > 3) or (
                    line.endswith(":") and len(line) < 40
                ):
                    if in_header:
                        # Name line
                        pdf.set_font("Helvetica", style="B", size=16)
                        pdf.cell(0, 10, self._safe_latin(line), ln=True, align="C")
                        in_header = False
                    else:
                        pdf.set_font("Helvetica", style="B", size=12)
                        pdf.ln(2)
                        pdf.cell(0, 8, self._safe_latin(line), ln=True)
                        pdf.line(pdf.get_x(), pdf.get_y(), 190, pdf.get_y())
                elif line.startswith("•") or line.startswith("-"):
                    pdf.set_font("Helvetica", size=10)
                    pdf.cell(5)
                    pdf.multi_cell(0, 6, self._safe_latin(line))
                else:
                    if in_header:
                        pdf.set_font("Helvetica", size=10)
                        pdf.cell(0, 6, self._safe_latin(line), ln=True, align="C")
                    else:
                        pdf.set_font("Helvetica", size=10)
                        pdf.multi_cell(0, 6, self._safe_latin(line))

            filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = self.output_dir / filename
            pdf.output(str(file_path))
            logger.info("Generated PDF: %s", file_path)
            return str(file_path)

        except Exception as exc:
            logger.error("PDF generation failed: %s", exc)
            raise

    # ──────────────────────────────────────────────────────────────
    # Helpers
    # ──────────────────────────────────────────────────────────────

    def _parse_sections(self, text: str) -> list[tuple[str, str]]:
        """Parse text into (type, content) tuples for DOCX formatting."""
        result: list[tuple[str, str]] = []
        lines = text.split("\n")

        is_first = True
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if is_first:
                result.append(("name", line_stripped))
                is_first = False
            elif line_stripped.endswith(":") and len(line_stripped) < 40:
                result.append(("header", line_stripped.rstrip(":")))
            elif line_stripped.startswith(("•", "-", "*")):
                result.append(("bullet", line_stripped.lstrip("•-* ")))
            elif "@" in line_stripped or re.search(r"\d{3}[-.\s]\d{3}", line_stripped):
                result.append(("contact", line_stripped))
            else:
                result.append(("paragraph", line_stripped))

        return result

    def _safe_latin(self, text: str) -> str:
        """Convert text to latin-1 safe string for FPDF."""
        return text.encode("latin-1", errors="replace").decode("latin-1")
